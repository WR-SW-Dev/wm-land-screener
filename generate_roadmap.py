"""Generate West Michigan Land Tool Roadmap as a Word document."""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT = "/Users/sadiebrooks/AI Development/MLS Tool Experiment/West_Michigan_Land_Tool_Roadmap.docx"

doc = Document()

# ── Page setup (US Letter, 1" margins) ──────────────────────────────────────
section = doc.sections[0]
section.page_width  = int(8.5 * 914400)
section.page_height = int(11  * 914400)
for attr in ("left_margin", "right_margin", "top_margin", "bottom_margin"):
    setattr(section, attr, Inches(1))

# ── Style helpers ─────────────────────────────────────────────────────────────
def set_heading_style(paragraph, level, text, color=None):
    paragraph.style = f"Heading {level}"
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run(text)
    if not paragraph.runs:
        run = paragraph.add_run(text)
    run.font.name = "Arial"
    if color:
        run.font.color.rgb = RGBColor(*color)
    return paragraph

def h1(text):
    p = doc.add_paragraph()
    p.style = "Heading 1"
    r = p.add_run(text)
    r.font.name = "Arial"
    r.font.size = Pt(16)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    return p

def h2(text):
    p = doc.add_paragraph()
    p.style = "Heading 2"
    r = p.add_run(text)
    r.font.name = "Arial"
    r.font.size = Pt(13)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    return p

def body(text, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Arial"
    r.font.size = Pt(11)
    r.font.bold = bold
    p.paragraph_format.space_after = Pt(4)
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(text)
    r.font.name = "Arial"
    r.font.size = Pt(11)
    p.paragraph_format.left_indent = Inches(0.25 + level * 0.25)
    p.paragraph_format.space_after = Pt(2)
    return p

def add_table(headers, rows, col_widths):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = "Table Grid"
    # Header row
    hdr = tbl.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.width = Inches(col_widths[i])
        p = cell.paragraphs[0]
        p.clear()
        r = p.add_run(h)
        r.font.name = "Arial"
        r.font.size = Pt(10)
        r.font.bold = True
        # header shading
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "1F497D")
        tcPr.append(shd)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    # Data rows
    for ri, row_data in enumerate(rows):
        row = tbl.rows[ri + 1]
        fill = "DEEAF1" if ri % 2 == 0 else "FFFFFF"
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            cell.width = Inches(col_widths[ci])
            p = cell.paragraphs[0]
            p.clear()
            r = p.add_run(val)
            r.font.name = "Arial"
            r.font.size = Pt(10)
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), fill)
            tcPr.append(shd)
    doc.add_paragraph()  # spacer after table

# ── Title ────────────────────────────────────────────────────────────────────
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
tr = title_p.add_run("West Michigan Vacant Land Tool")
tr.font.name = "Arial"
tr.font.size = Pt(24)
tr.font.bold = True
tr.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = sub_p.add_run("Build Roadmap")
sr.font.name = "Arial"
sr.font.size = Pt(16)
sr.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
sub_p.paragraph_format.space_after = Pt(24)

doc.add_paragraph()

# ── 1. Project Overview ───────────────────────────────────────────────────────
h1("1. Project Overview")
body(
    "A web application that screens and analyzes vacant land across Ottawa, Muskegon, and "
    "Allegan counties in West Michigan. The tool cross-references MLS listings, county "
    "parcel/GIS data, zoning layers, and utility service areas to identify feasible "
    "development sites and estimate potential unit density."
)
body("Starting scope: Grand Haven, Holland, and Muskegon.")
body("Long-term goal: Airtable integration for deal pipeline tracking.")

# ── 2. Tech Stack ─────────────────────────────────────────────────────────────
h1("2. Tech Stack")

add_table(
    headers=["Layer", "Technology", "Notes"],
    col_widths=[1.5, 2.5, 3.5],
    rows=[
        ["Backend",      "Python / FastAPI",                      "API layer and data processing"],
        ["Data/Spatial", "GeoPandas, Shapely, DuckDB + spatial",  "Parcel filtering and GIS overlays"],
        ["Frontend",     "Streamlit (MVP) → React (v2)",          "Interactive map and filter UI"],
        ["MLS",          "RESO Web API",                          "Via GRAR or WLAR OAuth credentials"],
        ["Maps",         "Folium / Leaflet.js",                   "Interactive parcel map"],
        ["Storage",      "DuckDB (MVP) → PostgreSQL + PostGIS",   "Upgrade when multi-user needed"],
        ["Future",       "Airtable API",                          "Deal pipeline sync (Phase 5)"],
    ]
)

# ── 3. Data Sources ───────────────────────────────────────────────────────────
h1("3. Data Sources")

add_table(
    headers=["Source", "Type", "Access Method", "Status"],
    col_widths=[2.0, 1.8, 2.2, 1.5],
    rows=[
        ["MLS (GRAR / WLAR)",         "Active listings + parcel match", "RESO Web API (OAuth)",       "Credentials needed"],
        ["Ottawa County GIS",          "Parcel data, ownership",         "Public download / WMS",      "Free"],
        ["Muskegon County GIS",        "Parcel data, ownership",         "Public download / WMS",      "Free"],
        ["Allegan County GIS",         "Parcel data, ownership",         "Public download / WMS",      "Free"],
        ["Municipal zoning layers",    "Allowed uses, density limits",   "GIS download / PDF",         "Varies by city"],
        ["FEMA Flood Map",             "Floodplain boundaries",          "Public WMS",                 "Free"],
        ["EGLE Wetlands",              "Wetland boundaries",             "Public WMS",                 "Free"],
        ["Water / Sewer districts",    "Utility service areas",          "County / utility GIS",       "Free (some manual)"],
        ["Michigan Statewide Parcel",  "Backup parcel data",             "MCGI download",              "Free"],
    ]
)

# ── 4. Feasibility Scoring ────────────────────────────────────────────────────
h1("4. Feasibility Scoring Logic")

h2("Hard Filters (parcel disqualified if any fail)")
bullet("Not in sewer service area")
bullet("Floodplain coverage > 25%")
bullet("Current use classification is not vacant / agricultural / unimproved")
bullet("Below minimum size threshold (default: 0.5 acres)")

h2("Soft Scoring (0–100)")
bullet("Zoning density — higher max units/acre scores higher")
bullet("Parcel size — larger scores higher up to diminishing returns (~5 acres)")
bullet("Wetland coverage — penalized above 10%")
bullet("Distance to utilities — score decreases as distance increases")
bullet("MLS listing status — on-market flagged; off-market flagged as opportunity")
bullet("Tax status — delinquent taxes or land bank ownership flagged as opportunity")

h2("Density Estimate")
bullet("Pull max units/acre from zoning lookup table for the parcel's zoning district")
bullet("Net developable area = gross acres minus floodplain and wetland overlap")
bullet("Output: estimated unit range — conservative (70% of max) and optimistic (100% of max)")

# ── 5. Build Phases ───────────────────────────────────────────────────────────
h1("5. Build Phases")

phases = [
    (
        "Phase 1 — Foundation (Weeks 1–3)",
        "Working data pipeline for one city (Grand Haven)",
        [
            "Set up project repo and Python environment (venv, requirements.txt)",
            "Download Ottawa County parcel data (shapefile / GeoJSON)",
            "Download Grand Haven zoning GIS layer",
            "Pull FEMA floodplain and EGLE wetland layers via public WMS",
            "Build parcel filtering script (vacant use codes, size threshold)",
            "Overlay zoning, floodplain, and wetlands onto parcels with GeoPandas",
            "Calculate net developable area and density estimate per parcel",
            "Output: CSV of qualifying parcels with feasibility scores",
        ]
    ),
    (
        "Phase 2 — MLS Integration (Weeks 3–5)",
        "Match MLS listings to qualifying parcels",
        [
            "Obtain RESO Web API credentials from GRAR / WLAR",
            "Build MLS API client (OAuth token handling, pagination, field mapping)",
            "Filter for vacant land / lot listings in target geographies",
            "Match MLS listings to parcel data by APN or address geocoding",
            "Flag on-market parcels vs. off-market opportunities",
            "Append listing price and days-on-market to parcel output",
        ]
    ),
    (
        "Phase 3 — Web App MVP (Weeks 5–8)",
        "Streamlit app with interactive map and parcel table",
        [
            "Build Streamlit UI with sidebar filters (county, city, min acres, min density)",
            "Interactive Folium map showing qualifying parcels color-coded by score",
            "Parcel detail panel: click a parcel to see zoning, density estimate, MLS status, score breakdown",
            "Export to CSV button",
            "Deploy locally; optionally push to Streamlit Cloud or a VPS",
        ]
    ),
    (
        "Phase 4 — Expand Coverage (Weeks 8–12)",
        "Add Holland and Muskegon; add utility layer",
        [
            "Repeat Phase 1 data pipeline for Holland (Ottawa Co.) and Muskegon (Muskegon Co.)",
            "Source water/sewer service area GIS files for each city",
            "Add utility overlay to feasibility hard filter",
            "Build full zoning lookup table for all three cities (allowed uses + max density by district)",
            "Refine scoring weights based on real-world testing with known sites",
        ]
    ),
    (
        "Phase 5 — Airtable Integration (Weeks 12–16)",
        "Push qualified parcels into Airtable deal tracker",
        [
            "Set up Airtable base with 'Deals' table schema (APN, address, acreage, units, score, etc.)",
            "Build Airtable API sync: create new records for new parcels, update existing",
            "Add 'Send to Airtable' button in Streamlit UI for manual pushes",
            "Synced fields: APN, address, acreage, estimated units, score, zoning, MLS status, listing price",
            "Optional: Airtable automation to notify team when new qualifying parcels are added",
        ]
    ),
]

for phase_title, phase_goal, tasks in phases:
    h2(phase_title)
    p = doc.add_paragraph()
    r = p.add_run("Goal: ")
    r.font.name = "Arial"
    r.font.size = Pt(11)
    r.font.bold = True
    r2 = p.add_run(phase_goal)
    r2.font.name = "Arial"
    r2.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(4)
    for task in tasks:
        bullet(task)

# ── 6. MLS API — Next Steps ───────────────────────────────────────────────────
h1("6. MLS API — Next Steps")
body(
    "Since you hold a license, API access is straightforward. Follow these steps to get connected:"
)

steps = [
    (
        "Step 1: Identify Your MLS",
        "West Michigan is served by the Greater Regional Alliance of REALTORS (GRAR) and the "
        "West Michigan Lakeshore Association of REALTORS (WLAR). Grand Haven and Holland fall "
        "under WLAR; Muskegon is covered by GRAR. You may need credentials for both."
    ),
    (
        "Step 2: Request RESO Web API Access",
        "Contact your MLS board and request 'RESO Web API' or 'Data Share' access for an "
        "internal business tool. You will likely need to sign a data access agreement. Noting "
        "that this is an internal tool (not a public IDX site) typically speeds up approval."
    ),
    (
        "Step 3: Receive OAuth Credentials",
        "Once approved, you will receive a client ID, client secret, and the API endpoint URL. "
        "Store these in environment variables — never in source code."
    ),
    (
        "Step 4: Test with Postman or Python",
        "Authenticate via OAuth and run a sample query. RESO Web API uses OData syntax. "
        "Example: GET /Property?$filter=PropertyType eq 'Land' and City eq 'Grand Haven'"
    ),
    (
        "Step 5: Build the Python Client",
        "Wrap API calls in a Python class that handles token refresh, pagination, and field "
        "mapping to your internal parcel schema. This becomes the MLS module in Phase 2."
    ),
]

for step_title, step_body in steps:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(2)
    r = p.add_run(step_title)
    r.font.name  = "Arial"
    r.font.size  = Pt(11)
    r.font.bold  = True
    p2 = doc.add_paragraph()
    r2 = p2.add_run(step_body)
    r2.font.name = "Arial"
    r2.font.size = Pt(11)
    p2.paragraph_format.left_indent = Inches(0.25)
    p2.paragraph_format.space_after = Pt(4)

# ── 7. Key Risks & Mitigations ────────────────────────────────────────────────
h1("7. Key Risks & Mitigations")

add_table(
    headers=["Risk", "Likelihood", "Mitigation"],
    col_widths=[2.5, 1.3, 3.7],
    rows=[
        ["MLS API approval takes time",
         "Medium",
         "Start Phases 1 and 3 in parallel; use ATTOM Data Solutions as a paid fallback"],
        ["Zoning data not available as GIS",
         "High",
         "Build manual zoning lookup table as fallback; scrape ordinance PDFs for density rules"],
        ["County parcel data quality varies",
         "Medium",
         "Cross-reference with Michigan statewide parcel layer (MCGI) for gap-filling"],
        ["Wetland / floodplain data misses small features",
         "Low–Medium",
         "Flag edge-case parcels for manual review rather than hard-disqualifying them"],
        ["Scope creep across 3 counties",
         "High",
         "Enforce strict phase gates — do not expand geography until current phase is stable"],
    ]
)

# ── 8. Future Enhancements ────────────────────────────────────────────────────
h1("8. Future Enhancements (Post-MVP)")

enhancements = [
    "Full Allegan County coverage",
    "Automated weekly refresh of parcel and MLS data",
    "School district overlay and walkability scoring",
    "Comparable sales analysis using MLS sold data",
    "Permit history lookup by parcel",
    "Per-parcel PDF report generation",
    "Multi-user login and saved search profiles",
    "Email / Slack alerts for new qualifying parcels",
]
for e in enhancements:
    bullet(e)

# ── Save ──────────────────────────────────────────────────────────────────────
doc.save(OUTPUT)
print(f"Saved: {OUTPUT}")
